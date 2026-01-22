"""DOCX report generator for search strategies."""

from datetime import datetime
from io import BytesIO
from typing import Optional

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DOCXGenerator:
    """Generate PRISMA-compliant DOCX reports for search strategies."""

    def __init__(self):
        """Initialize DOCX generator."""
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for DOCX generation. "
                "Install with: pip install python-docx"
            )

    def generate_search_report(
        self,
        research_question: str = "",
        strategies: dict[str, str] = None,
        dedup_result=None,
        author: str = "",
        date_searched: Optional[datetime] = None,
        filters_applied: Optional[list[str]] = None,
    ) -> bytes:
        """
        Generate a PRISMA-compliant search strategy report.

        Args:
            research_question: The research question
            strategies: Dictionary mapping database name to strategy
            dedup_result: DeduplicationResult object
            author: Author name
            date_searched: Date searches were conducted
            filters_applied: List of filters applied

        Returns:
            DOCX file as bytes
        """
        doc = Document()

        # Title
        title = doc.add_heading("Search Strategy Report", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        if author:
            doc.add_paragraph(f"Author: {author}")
        if date_searched:
            doc.add_paragraph(f"Date Searched: {date_searched.strftime('%Y-%m-%d')}")

        doc.add_paragraph()  # Spacer

        # Research Question
        if research_question:
            doc.add_heading("Research Question", level=1)
            doc.add_paragraph(research_question)

        # Search Strategies
        if strategies:
            doc.add_heading("Search Strategies", level=1)

            for db_name, strategy in strategies.items():
                doc.add_heading(db_name, level=2)

                # Add strategy as monospace
                p = doc.add_paragraph()
                run = p.add_run(strategy)
                run.font.name = "Courier New"
                run.font.size = Pt(9)

                doc.add_paragraph()  # Spacer

        # Filters Applied
        if filters_applied:
            doc.add_heading("Filters Applied", level=1)
            for filter_text in filters_applied:
                doc.add_paragraph(f"• {filter_text}", style="List Bullet")

        # Results Summary
        if dedup_result:
            doc.add_heading("Search Results", level=1)

            # Records per source table
            doc.add_heading("Records by Database", level=2)

            table = doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"

            header_cells = table.rows[0].cells
            header_cells[0].text = "Database"
            header_cells[1].text = "Records"

            for source, count in dedup_result.records_per_source.items():
                row_cells = table.add_row().cells
                row_cells[0].text = source
                row_cells[1].text = str(count)

            # Total row
            row_cells = table.add_row().cells
            row_cells[0].text = "Total"
            row_cells[0].paragraphs[0].runs[0].bold = True
            row_cells[1].text = str(dedup_result.total_records)
            row_cells[1].paragraphs[0].runs[0].bold = True

            doc.add_paragraph()  # Spacer

            # Deduplication summary
            doc.add_heading("Deduplication Summary", level=2)

            dedup_table = doc.add_table(rows=5, cols=2)
            dedup_table.style = "Table Grid"

            data = [
                ("Total Records", str(dedup_result.total_records)),
                ("Unique Records", str(dedup_result.unique_records)),
                ("Duplicates Removed", str(dedup_result.duplicate_count)),
                ("DOI Matches", str(dedup_result.doi_duplicates)),
                ("Title/Author Matches", str(
                    dedup_result.title_fuzzy_duplicates +
                    dedup_result.title_author_year_duplicates
                )),
            ]

            for i, (label, value) in enumerate(data):
                dedup_table.rows[i].cells[0].text = label
                dedup_table.rows[i].cells[1].text = value

        # PRISMA Flow Note
        doc.add_paragraph()
        doc.add_heading("PRISMA 2020 Compliance Note", level=1)

        prisma_note = (
            "This search strategy report is designed to support PRISMA 2020 "
            "guidelines for reporting systematic reviews. The search strategies "
            "should be reported in full for each database searched, including "
            "the date of search and any limits or filters applied."
        )
        doc.add_paragraph(prisma_note)

        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer.getvalue()

    def generate_full_report(
        self,
        project_name: str,
        research_question: str,
        pico_analysis: Optional[dict] = None,
        concept_blocks: Optional[list] = None,
        strategies: dict[str, str] = None,
        dedup_result=None,
        screening_summary: Optional[dict] = None,
    ) -> bytes:
        """
        Generate a comprehensive systematic review report.

        Args:
            project_name: Name of the project
            research_question: The research question
            pico_analysis: PICO analysis dictionary
            concept_blocks: List of concept blocks
            strategies: Dictionary of search strategies
            dedup_result: Deduplication results
            screening_summary: Summary of screening results

        Returns:
            DOCX file as bytes
        """
        doc = Document()

        # Title Page
        doc.add_heading(project_name, 0)
        doc.add_paragraph("Systematic Review Search Strategy Report")
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d')}")

        doc.add_page_break()

        # Table of Contents placeholder
        doc.add_heading("Table of Contents", level=1)
        doc.add_paragraph("[Table of contents - update after editing]")
        doc.add_page_break()

        # 1. Introduction
        doc.add_heading("1. Research Question", level=1)
        if research_question:
            doc.add_paragraph(research_question)

        # 2. PICO Analysis
        if pico_analysis:
            doc.add_heading("2. PICO Framework Analysis", level=1)

            for element in ["population", "intervention", "comparison", "outcome"]:
                if element in pico_analysis and pico_analysis[element]:
                    doc.add_heading(element.capitalize(), level=2)
                    data = pico_analysis[element]

                    doc.add_paragraph(f"Label: {data.get('label', '')}")

                    if data.get("primary_terms"):
                        doc.add_paragraph(
                            f"Primary Terms: {', '.join(data['primary_terms'])}"
                        )

                    if data.get("mesh_terms"):
                        doc.add_paragraph(
                            f"MeSH Terms: {', '.join(data['mesh_terms'])}"
                        )

        # 3. Concept Blocks
        if concept_blocks:
            doc.add_heading("3. Search Concepts", level=1)

            for block in concept_blocks:
                doc.add_heading(block.name, level=2)
                elem = block.pico_element

                if elem.primary_terms:
                    doc.add_paragraph(f"Primary Terms: {', '.join(elem.primary_terms)}")
                if elem.synonyms:
                    doc.add_paragraph(f"Synonyms: {', '.join(elem.synonyms)}")
                if elem.mesh_terms:
                    doc.add_paragraph(f"MeSH Terms: {', '.join(elem.mesh_terms)}")

        # 4. Search Strategies
        if strategies:
            doc.add_heading("4. Database Search Strategies", level=1)

            for db_name, strategy in strategies.items():
                doc.add_heading(f"4.{list(strategies.keys()).index(db_name)+1}. {db_name}", level=2)

                # Strategy text
                p = doc.add_paragraph()
                run = p.add_run(strategy)
                run.font.name = "Courier New"
                run.font.size = Pt(9)

        # 5. Search Results
        doc.add_heading("5. Search Results", level=1)

        if dedup_result:
            # Results table
            table = doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"

            header = table.rows[0].cells
            header[0].text = "Database"
            header[1].text = "Records Retrieved"

            for source, count in dedup_result.records_per_source.items():
                row = table.add_row().cells
                row[0].text = source
                row[1].text = str(count)

            # Totals
            row = table.add_row().cells
            row[0].text = "Total before deduplication"
            row[1].text = str(dedup_result.total_records)

            row = table.add_row().cells
            row[0].text = "Duplicates removed"
            row[1].text = str(dedup_result.duplicate_count)

            row = table.add_row().cells
            row[0].text = "Total after deduplication"
            row[0].paragraphs[0].runs[0].bold = True
            row[1].text = str(dedup_result.unique_records)
            row[1].paragraphs[0].runs[0].bold = True

        # 6. Screening Summary (if available)
        if screening_summary:
            doc.add_heading("6. Screening Summary", level=1)

            for key, value in screening_summary.items():
                doc.add_paragraph(f"{key}: {value}")

        # Appendix
        doc.add_page_break()
        doc.add_heading("Appendix: PRISMA 2020 Checklist Items", level=1)

        checklist_items = [
            "METHODS - Eligibility criteria (Item 5)",
            "METHODS - Information sources (Item 6)",
            "METHODS - Search strategy (Item 7)",
            "RESULTS - Study selection (Item 16)",
        ]

        for item in checklist_items:
            doc.add_paragraph(f"☐ {item}", style="List Bullet")

        # Save
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer.getvalue()
